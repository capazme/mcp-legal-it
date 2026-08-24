"""Generazione DOCX pronti-firma per il recupero crediti seriale: procura alle liti
ex art. 83 c.p.c. e lettera di quotazione compensi (D.M. 55/2014 agg. D.M. 147/2022)
per procedimento monitorio, esecuzione forzata e opposizione a decreto ingiuntivo."""

import json
import os
import re
import tempfile
import uuid
from datetime import date
from decimal import ROUND_HALF_UP, Decimal
from pathlib import Path

from src.server import mcp
from src.lib._data import sourced

_DATA = Path(__file__).resolve().parent.parent / "data"

with open(_DATA / "parametri_forensi.json") as f:
    _PARAMETRI = json.load(f)

with open(_DATA / "contributo_unificato.json") as f:
    _CU = json.load(f)

_OUTPUT_DIR = os.path.join(tempfile.gettempdir(), "mcp-legal-it")

# ---------------------------------------------------------------------------
# Tabelle D.M. 55/2014 (agg. D.M. 147/2022) — procedimenti monitori, fase unica.
# Coppie (soglia_scaglione, minimo, medio) dalla tabella ministeriale: il minimo
# pubblicato e' il medio ridotto del 50% ex art. 4, co. 1 (arrotondato: per i
# primi scaglioni il medio e' dispari, quindi minimo x 2 NON restituisce il medio).
# ---------------------------------------------------------------------------
_MONITORI_FASE_UNICA = [
    (5200, Decimal("237"), Decimal("473")),
    (26000, Decimal("284"), Decimal("567")),
    (52000, Decimal("685"), Decimal("1370")),
    (260000, Decimal("1121"), Decimal("2242")),
    (520000, Decimal("2197"), Decimal("4394")),
]

# Compensi esecuzione forzata di default (D.M. 55/2014 agg. 147/2022, esecuzioni
# mobiliari, scaglione fino a € 5.200, valori minimi). Oltre questo scaglione o a
# livello diverso i compensi vanno passati esplicitamente dal chiamante.
_ESECUZIONE_DEFAULT_MAX_VALORE = 5200
_ESECUZIONE_DEFAULT_INTRODUTTIVA = 166.0
_ESECUZIONE_DEFAULT_TRATTAZIONE = 284.0

_FASI_OPPOSIZIONE = ["studio", "introduttiva", "istruttoria", "decisionale"]
_ETICHETTE_FASI = {
    "studio": "Fase di studio",
    "introduttiva": "Fase introduttiva del giudizio",
    "istruttoria": "Fase istruttoria/di trattazione",
    "decisionale": "Fase decisionale",
}

_MESI = [
    "gennaio", "febbraio", "marzo", "aprile", "maggio", "giugno",
    "luglio", "agosto", "settembre", "ottobre", "novembre", "dicembre",
]


def _d2(value) -> Decimal:
    return Decimal(str(value)).quantize(Decimal("0.01"), ROUND_HALF_UP)


def _eur(value) -> str:
    testo = f"{_d2(value):,.2f}"
    return "€ " + testo.replace(",", "X").replace(".", ",").replace("X", ".")


def _data_lettere(data_doc: str) -> str:
    """Converte GG/MM/AAAA in forma estesa italiana; testo libero passa invariato."""
    if not data_doc:
        oggi = date.today()
        return f"{oggi.day} {_MESI[oggi.month - 1]} {oggi.year}"
    m = re.fullmatch(r"(\d{1,2})/(\d{1,2})/(\d{4})", data_doc.strip())
    if m:
        giorno, mese, anno = int(m.group(1)), int(m.group(2)), int(m.group(3))
        if 1 <= mese <= 12:
            return f"{giorno} {_MESI[mese - 1]} {anno}"
    return data_doc.strip()


def _sanitize_filename(name: str) -> str:
    name = name.lower().replace(" ", "_")
    name = re.sub(r"[^a-z0-9_\-]", "", name)
    return name[:50] or "documento"


def _salva(doc, prefisso: str) -> str:
    os.makedirs(_OUTPUT_DIR, exist_ok=True)
    filename = f"{_sanitize_filename(prefisso)}_{uuid.uuid4().hex[:8]}.docx"
    filepath = os.path.join(_OUTPUT_DIR, filename)
    doc.save(filepath)
    size_kb = round(os.path.getsize(filepath) / 1024, 1)
    return f"File salvato: {filepath} ({size_kb} KB)"


def _monitorio_tabellare(valore: float, livello: str) -> Decimal | None:
    for soglia, minimo, medio in _MONITORI_FASE_UNICA:
        if valore <= soglia:
            return medio if livello == "medi" else minimo
    return None


def _cu_monitorio(valore: float) -> Decimal:
    """Contributo unificato del procedimento monitorio dalla tabella canonica
    del repo (DPR 115/2002, dimezzato ex art. 13, co. 3)."""
    for scaglione in _CU["civile"]["procedimento_monitorio"]["scaglioni"]:
        if scaglione.get("oltre") or valore <= scaglione["fino_a"]:
            return _d2(scaglione["importo"])
    return _d2(_CU["civile"]["procedimento_monitorio"]["scaglioni"][-1]["importo"])


def _prospetto_importi(tabellare: Decimal, aumento_pct30: bool) -> dict:
    """Catena di calcolo validata: +30% PCT, SG 15%, CPA 4%, IVA 22%, RA 20%."""
    aumento = _d2(tabellare * Decimal("0.30")) if aumento_pct30 else Decimal("0.00")
    totale = _d2(tabellare + aumento)
    spese_generali = _d2(totale * Decimal("0.15"))
    base_cpa = totale + spese_generali
    cpa = _d2(base_cpa * Decimal("0.04"))
    imponibile = _d2(totale + spese_generali + cpa)
    iva = _d2(imponibile * Decimal("0.22"))
    liquidabile = _d2(imponibile + iva)
    ritenuta = _d2(base_cpa * Decimal("0.20"))
    totale_documento = _d2(liquidabile - ritenuta)
    return {
        "tabellare": _d2(tabellare),
        "aumento": aumento,
        "totale": totale,
        "spese_generali": spese_generali,
        "cpa": cpa,
        "imponibile": imponibile,
        "iva": iva,
        "liquidabile": liquidabile,
        "ritenuta": ritenuta,
        "totale_documento": totale_documento,
    }


@mcp.tool(tags={"atti", "credito", "giudiziario"})
def genera_procura_liti_docx(
    mandante_denominazione: str,
    mandante_sede: str,
    mandante_cf_piva: str,
    firmatario_nome: str,
    firmatario_cf: str,
    controparte: str,
    difensori: list[dict],
    domicilio_studio: str,
    pec: str,
    firmatario_qualifica: str = "legale rappresentante",
    fax: str = "",
    luogo: str = "Milano",
    data_documento: str = "",
) -> str:
    """Genera la procura alle liti ex art. 83, co. 3, c.p.c. in DOCX pronto-firma (una pagina).

    Produce il modello completo usato nel recupero crediti seriale: conferimento a uno o
    piu' difensori (congiuntamente e disgiuntamente), elezione di domicilio, dichiarazioni
    su mediazione (art. 4 D.Lgs. 28/2010), negoziazione assistita (D.L. 132/2014),
    preventivo e polizza, consenso privacy, apposizione in calce ex art. 18 D.M. 44/2011,
    blocco firma del mandante e autentica dei difensori. Formattazione: Times New Roman 11,
    impaginata su una sola pagina. Usare quando serve la procura in Word da far firmare;
    per un semplice testo base usare invece procura_alle_liti().
    Vigenza: art. 83 c.p.c.; art. 18, co. 5, D.M. 44/2011. Precisione: ESATTO (testo modello).

    Args:
        mandante_denominazione: Denominazione della parte che conferisce la procura (es. "Esempio S.r.l.")
        mandante_sede: Sede legale completa (es. "via Roma n. 1 – 20100 Milano (MI)")
        mandante_cf_piva: Codice fiscale e/o Partita IVA del mandante
        firmatario_nome: Nome e cognome della persona fisica che firma
        firmatario_cf: Codice fiscale del firmatario
        controparte: Clausola identificativa della controparte, VERBATIM come da atti (denominazione, legale rappr., sede, C.F./P.IVA)
        difensori: Lista di difensori, ciascuno come oggetto {"nome": "Giulia Bianchi", "cf": "BNCGLI80A41F205Y"}
        domicilio_studio: Indirizzo dello studio per l'elezione di domicilio (es. "corso Esempio n. 10 – 20100 Milano (MI)")
        pec: PEC del/dei difensori come da testo (es. "g.bianchi@pec.esempio.it e p.verdi@pec.esempio.it")
        firmatario_qualifica: Qualifica del firmatario (es. "presidente del consiglio di amministrazione e legale rappresentante")
        fax: Numero fax dello studio (opzionale)
        luogo: Luogo di sottoscrizione (default Milano)
        data_documento: Data in formato GG/MM/AAAA (convertita in forma estesa) o testo libero; vuota = data odierna
    """
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Cm, Pt
    except ImportError:
        return "Errore: python-docx non installato. Eseguire: pip install python-docx"

    if not mandante_denominazione.strip():
        return "Errore: mandante_denominazione è obbligatorio."
    if not controparte.strip():
        return "Errore: controparte è obbligatoria (clausola identificativa verbatim)."
    if not difensori:
        return "Errore: indicare almeno un difensore ({'nome': ..., 'cf': ...})."
    for d in difensori:
        if not isinstance(d, dict) or not d.get("nome"):
            return "Errore: ogni difensore deve essere un oggetto {'nome': ..., 'cf': ...}."

    doc = Document()
    stile = doc.styles["Normal"]
    stile.font.name = "Times New Roman"
    stile.font.size = Pt(11)
    for sezione in doc.sections:
        sezione.top_margin = Cm(2.0)
        sezione.bottom_margin = Cm(1.5)
        sezione.left_margin = Cm(2.0)
        sezione.right_margin = Cm(2.0)

    def par(testo="", bold=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY, hang=False):
        p = doc.add_paragraph()
        p.alignment = align
        pf = p.paragraph_format
        pf.space_before = Pt(0)
        pf.space_after = Pt(2)
        pf.line_spacing = 1.0
        if hang:
            pf.left_indent = Cm(0.5)
            pf.first_line_indent = Cm(-0.5)
        if testo:
            r = p.add_run(testo)
            r.bold = bold
        return p

    par("PROCURA ALLE LITI", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    par("RILASCIATA AI SENSI DELL’ART. 83, COMMA 3, C.P.C.", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)

    par(
        f"Il sottoscritto {firmatario_nome} (Cod. Fisc. {firmatario_cf}), nella sua qualità di "
        f"{firmatario_qualifica} di {mandante_denominazione}, con sede legale in {mandante_sede}, "
        f"Cod. Fisc. e Partita IVA {mandante_cf_piva},"
    )
    par("CONFERISCE PROCURA AD LITEM", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)

    elenco = " e ".join(
        f"all’avv. {d['nome']}" + (f" (Cod. Fisc. {d['cf']})" if d.get("cf") else "")
        for d in difensori
    )
    congiunzione = "congiuntamente e disgiuntamente tra loro, " if len(difensori) > 1 else ""
    par(
        f"{elenco}, {congiunzione}per rappresentare e difendere la predetta parte nel procedimento "
        f"nei confronti di {controparte}, ed in ogni successiva fase e grado, compresa quella di appello, "
        "reclamo, opposizione ed esecuzione, conferendo loro all’uopo ogni più ampia facoltà di legge, "
        "ivi comprese, a titolo meramente esemplificativo e non esaustivo, la facoltà di transigere, "
        "conciliare, incassare, rinunciare agli atti, farsi rappresentare, assistere e sostituire, "
        "indicare domiciliatari, riassumere la causa, proseguirla, chiamare terzi in causa, deferire "
        "giuramento, proporre domande riconvenzionali ed azioni cautelari di qualsiasi genere, dando "
        "sin d’ora per rato e valido il loro operato senza bisogno di alcuna ratifica espressa, e con "
        "espresso potere di dichiararsi antistatari."
    )
    recapiti = f"fax {fax}, PEC {pec}" if fax else f"PEC {pec}"
    plurale = "i nominati difensori" if len(difensori) > 1 else "il nominato difensore"
    par(f"Elegge domicilio presso {plurale} con studio in {domicilio_studio}, {recapiti}.")

    par("DICHIARA", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    dichiarazioni = [
        "di voler ricevere le comunicazioni, le notifiche e gli avvisi relativi al presente procedimento al domicilio eletto;",
        "di essere stato informato, ai sensi dell’art. 4, co. 3, D.Lgs. n. 28/2010 e ss.mm.ii., della possibilità di ricorrere al procedimento di mediazione ivi previsto e dei benefici fiscali di cui agli artt. 17 e 20 del medesimo decreto, nonché dei casi in cui l’esperimento del procedimento di mediazione è condizione di procedibilità della domanda giudiziale;",
        "di essere stato informato, ai sensi dell’art. 2, co. 7, D.L. n. 132/2014, convertito in L. n. 162/2014, della possibilità di ricorrere alla convenzione di negoziazione assistita da uno o più avvocati disciplinata dagli artt. 2 e ss. del suddetto decreto legge, nonché dei casi di cui all’art. 3 del suddetto decreto in cui l’esperimento di tale procedimento è condizione di procedibilità della domanda giudiziale;",
        "di essere stato reso edotto circa i rischi del contenzioso e il grado di complessità dell’incarico che con la presente conferisce, delle caratteristiche e dell’importanza dell’incarico, delle attività da espletare, delle iniziative da intraprendere, delle ipotesi di soluzione e della prevedibile durata del processo;",
        "di avere ricevuto tutte le informazioni utili circa gli oneri ipotizzabili dal momento del conferimento sino alla conclusione dell’incarico, nonché di aver ricevuto ed accettato un preventivo scritto relativo alla prevedibile misura dei costi della prestazione, con distinzione analitica delle voci di costo tra oneri, anche fiscali e previdenziali, spese, anche forfettarie, e compenso professionale;",
        "di essere stato reso edotto degli estremi della polizza assicurativa professionale dei suoi difensori e della facoltà di recedere dall’incarico professionale costituendo giusta causa il mancato pagamento degli oneri così come convenuti e formalizzati nelle proforma da trasmettersi;",
        "ai sensi e per gli effetti di cui al D.Lgs. n. 196/2003 e del Regolamento UE n. 679/2016 e ss.mm.ii., di essere stato informato che i dati personali, anche sensibili, verranno utilizzati per le finalità inerenti al presente mandato, autorizzando sin d’ora il rispettivo trattamento.",
    ]
    for voce in dichiarazioni:
        par("-  " + voce, hang=True)

    par(
        "La presente procura alle liti è da intendersi apposta in calce all’atto, anche ai sensi "
        "dell’art. 18, co. 5, D.M. Giustizia n. 44/2011, come sostituito dal D.M. Giustizia n. 48/2013."
    )
    par(f"{luogo}, {_data_lettere(data_documento)}", align=WD_ALIGN_PARAGRAPH.LEFT)
    par("Firma", align=WD_ALIGN_PARAGRAPH.LEFT)
    par(f"{firmatario_nome}\t\t_________________________________", align=WD_ALIGN_PARAGRAPH.LEFT)
    par("È vera e autentica", align=WD_ALIGN_PARAGRAPH.LEFT)
    for d in difensori:
        par(f"Avv. {d['nome']}\t\t_________________________________", align=WD_ALIGN_PARAGRAPH.LEFT)

    controparte_breve = controparte.split(",")[0].strip()
    return _salva(doc, f"procura_{mandante_denominazione}_{controparte_breve}")


@mcp.tool(tags={"atti", "credito", "parcelle_avv"})
@sourced("contributo_unificato", "parametri_forensi")
def genera_quotazione_docx(
    tipo: str,
    valore_causa: float,
    debitore: str,
    cliente_denominazione: str,
    cliente_indirizzo: str,
    difensori: list[str],
    livello: str = "minimi",
    accettazione_denominazione: str = "",
    luogo: str = "Milano",
    data_documento: str = "",
    contributo_unificato: float = -1,
    compenso_fase_introduttiva: float = 166,
    compenso_fase_trattazione: float = 284,
) -> str:
    """Genera la lettera di quotazione compensi (D.M. 55/2014 agg. D.M. 147/2022) in DOCX.

    Lettera al cliente con prospetto di liquidazione, oneri accessori (contributo unificato,
    marca da bollo), nota sull'imposta di registro, firma dei difensori e blocco di
    accettazione del cliente in calce. Tre tipi: 'monitorio' (fase unica + aumento 30% PCT
    ex art. 4, co. 1-bis, con ritenuta d'acconto), 'esecuzione' (fase introduttiva +
    trattazione/conclusiva, oneri pignoramento), 'opposizione' (giudizio ordinario a fasi
    piene + aumento 30% PCT). Usare nel recupero crediti seriale insieme a
    genera_procura_liti_docx(); scegliere il tipo in base alla fase reale della posizione
    (decreto esecutivo -> esecuzione; decreto opposto -> opposizione).
    Vigenza: D.M. 55/2014 agg. D.M. 147/2022; CU ex DPR 115/2002 (monitorio: ridotto alla metà).
    Precisione: INDICATIVO per i compensi (valori tabellari); ESATTO per la catena SG 15%/CPA 4%/IVA 22%.

    Args:
        tipo: Tipo di quotazione: 'monitorio', 'esecuzione' o 'opposizione'
        valore_causa: Valore della causa/credito in euro (€)
        debitore: Denominazione del debitore come deve comparire in lettera (es. "Delta S.r.l.")
        cliente_denominazione: Denominazione del cliente destinatario (es. "Esempio S.r.l.")
        cliente_indirizzo: Indirizzo del cliente, righe separate da ';' o a capo (es. "via Roma n. 1; 20100 - Milano")
        difensori: Nomi dei difensori firmatari (es. ["Avv. Giulia Bianchi", "Avv. Paolo Verdi"])
        livello: 'minimi' o 'medi' (valori della tabella ministeriale; il minimo è il medio ridotto del 50% ex art. 4 D.M. 55/2014)
        accettazione_denominazione: Denominazione nel blocco di accettazione (default: cliente_denominazione)
        luogo: Luogo della lettera (default Milano)
        data_documento: Data GG/MM/AAAA (convertita in forma estesa) o testo libero; vuota = odierna
        contributo_unificato: CU in euro; -1 = calcolo automatico (monitorio: metà DPR 115/2002; esecuzione: € 139,00)
        compenso_fase_introduttiva: Solo tipo 'esecuzione': compenso fase introduttiva. Il default € 166,00 vale SOLO per valore causa fino a € 5.200 a livello minimi; oltre, o a livello 'medi', va passato il valore corretto della tabella esecuzioni
        compenso_fase_trattazione: Solo tipo 'esecuzione': compenso fase trattazione/conclusiva. Il default € 284,00 vale SOLO per valore causa fino a € 5.200 a livello minimi; oltre, o a livello 'medi', va passato il valore corretto
    """
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Cm, Pt
    except ImportError:
        return "Errore: python-docx non installato. Eseguire: pip install python-docx"

    tipo = tipo.strip().lower()
    if tipo not in ("monitorio", "esecuzione", "opposizione"):
        return f"Errore: tipo non valido: {tipo}. Usare: monitorio, esecuzione, opposizione"
    if livello not in ("minimi", "medi"):
        return f"Errore: livello non valido: {livello}. Usare: minimi, medi"
    if valore_causa <= 0:
        return "Errore: valore_causa deve essere positivo."
    if not difensori:
        return "Errore: indicare almeno un difensore firmatario."

    livello_label = "valori medi" if livello == "medi" else "valori minimi"
    fase_label = "valore medio" if livello == "medi" else "valore minimo"

    # ---- calcolo compensi per tipo ----
    righe_fasi: list[tuple[str, str]] = []
    if tipo == "monitorio":
        tabellare = _monitorio_tabellare(valore_causa, livello)
        if tabellare is None:
            return (
                "Errore: valore causa oltre € 520.000 non coperto dalla tabella "
                "procedimenti monitori inclusa nel tool."
            )
        importi = _prospetto_importi(tabellare, aumento_pct30=True)
        righe_fasi = [(f"Fase unica, {fase_label}:", _eur(tabellare))]
        oggetto = f"Quotazione giudiziaria procedimento monitorio {debitore}"
        corpo_1 = (
            "con la presente Vi trasmettiamo il prospetto di liquidazione dei compensi professionali "
            f"maturati in relazione al procedimento monitorio instaurato nei confronti di {debitore}"
        )
    elif tipo == "opposizione":
        liv_json = "medio" if livello == "medi" else "min"
        scaglione = None
        for s in _PARAMETRI["civile"]["scaglioni"]:
            if s.get("oltre") or valore_causa <= s["fino_a"]:
                scaglione = s
                break
        tabellare = Decimal("0")
        for fase in _FASI_OPPOSIZIONE:
            importo = Decimal(str(scaglione[fase][liv_json]))
            tabellare += importo
            righe_fasi.append((f"{_ETICHETTE_FASI[fase]}, {fase_label}:", _eur(importo)))
        importi = _prospetto_importi(tabellare, aumento_pct30=True)
        oggetto = f"Quotazione giudiziaria giudizio di opposizione a decreto ingiuntivo {debitore}"
        corpo_1 = (
            "con la presente Vi trasmettiamo il prospetto di liquidazione dei compensi professionali "
            "prevedibili in relazione al giudizio di opposizione a decreto ingiuntivo "
            f"(art. 645 c.p.c.) instaurato da {debitore}"
        )
    else:  # esecuzione
        if compenso_fase_introduttiva <= 0 or compenso_fase_trattazione <= 0:
            return "Errore: i compensi di fase dell'esecuzione devono essere positivi."
        default_compensi = (
            compenso_fase_introduttiva == _ESECUZIONE_DEFAULT_INTRODUTTIVA
            and compenso_fase_trattazione == _ESECUZIONE_DEFAULT_TRATTAZIONE
        )
        if default_compensi and livello == "medi":
            return (
                "Errore: i compensi di default dell'esecuzione sono i MINIMI dello "
                "scaglione fino a € 5.200: per una quotazione a valori medi passare "
                "compenso_fase_introduttiva e compenso_fase_trattazione della tabella esecuzioni."
            )
        if default_compensi and valore_causa > _ESECUZIONE_DEFAULT_MAX_VALORE:
            return (
                "Errore: per valore causa oltre € 5.200 i compensi di default "
                "dell'esecuzione (scaglione base) non sono applicabili: passare "
                "compenso_fase_introduttiva e compenso_fase_trattazione dello scaglione corretto."
            )
        intro = _d2(compenso_fase_introduttiva)
        tratt = _d2(compenso_fase_trattazione)
        tabellare = intro + tratt
        importi = _prospetto_importi(tabellare, aumento_pct30=False)
        righe_fasi = [
            (f"Fase introduttiva del giudizio, {fase_label}:", _eur(intro)),
            (f"Fase di trattazione e conclusiva, {fase_label}:", _eur(tratt)),
        ]
        oggetto = f"Quotazione giudiziaria esecuzione forzata {debitore}"
        corpo_1 = (
            "con la presente Vi trasmettiamo il prospetto di liquidazione dei compensi professionali "
            f"maturati in relazione all’atto di esecuzione forzata nei confronti di {debitore}"
        )

    # ---- documento ----
    doc = Document()
    stile = doc.styles["Normal"]
    stile.font.name = "Aptos"
    stile.font.size = Pt(12)
    for sezione in doc.sections:
        sezione.top_margin = Cm(2.5)
        sezione.bottom_margin = Cm(2.0)
        sezione.left_margin = Cm(2.0)
        sezione.right_margin = Cm(2.0)

    def par(testo="", bold=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY, dopo=6):
        p = doc.add_paragraph()
        p.alignment = align
        pf = p.paragraph_format
        pf.space_before = Pt(0)
        pf.space_after = Pt(dopo)
        pf.line_spacing = 1.0
        if testo:
            r = p.add_run(testo)
            r.bold = bold
        return p

    par("Spett.le Società", align=WD_ALIGN_PARAGRAPH.LEFT, dopo=0)
    par(cliente_denominazione, align=WD_ALIGN_PARAGRAPH.LEFT, dopo=0)
    for riga in re.split(r";|\n", cliente_indirizzo):
        if riga.strip():
            par(riga.strip(), align=WD_ALIGN_PARAGRAPH.LEFT, dopo=0)
    par("Inviata tramite Mail", align=WD_ALIGN_PARAGRAPH.LEFT, dopo=12)
    par(f"{luogo}, {_data_lettere(data_documento)}", align=WD_ALIGN_PARAGRAPH.RIGHT, dopo=12)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(12)
    p.paragraph_format.line_spacing = 1.0
    r = p.add_run("Oggetto: ")
    r.bold = True
    r = p.add_run(oggetto)
    r.bold = True

    par("Spettabile Società,", align=WD_ALIGN_PARAGRAPH.LEFT)
    par(corpo_1 + ("" if debitore.rstrip().endswith(".") else "."))
    par(
        "I compensi sono stati determinati in conformità alle tariffe di cui al D.M. 13 agosto 2022, "
        "n. 147, pubblicato nella Gazzetta Ufficiale n. 236 dell'8 ottobre 2022 ed in vigore dal "
        f"23 ottobre 2022, applicando i {livello_label} previsti dalle tabelle ministeriali."
    )
    par("Di seguito il dettaglio del prospetto:")
    par(f"Valore della causa: {_eur(valore_causa)}", bold=True, align=WD_ALIGN_PARAGRAPH.LEFT)

    righe: list[tuple[str, str]] = [("Fase", "Compenso")]
    righe += righe_fasi
    if tipo == "esecuzione":
        righe += [
            (f"Compenso tabellare ({livello_label})", _eur(importi["tabellare"])),
            ("Spese generali ( 15% sul compenso totale )", _eur(importi["spese_generali"])),
            ("Cassa Avvocati ( 4% )", _eur(importi["cpa"])),
            ("Totale imponibile", _eur(importi["imponibile"])),
            ("IVA 22% su Imponibile", _eur(importi["iva"])),
            ("IPOTESI DI COMPENSO LIQUIDABILE", _eur(importi["liquidabile"])),
        ]
    else:
        if tipo == "opposizione":
            righe.append((f"Compenso tabellare ({livello_label})", _eur(importi["tabellare"])))
        righe += [
            ("Aumento del 30% per tecniche informatiche PCT (art. 4, co. 1 bis)", _eur(importi["aumento"])),
            ("PROSPETTO FINALE", ""),
            ("Compenso tabellare", _eur(importi["tabellare"])),
            ("Totale variazioni in aumento", "+ " + _eur(importi["aumento"])),
            ("Compenso totale", _eur(importi["totale"])),
            ("Spese generali ( 15% sul compenso totale )", _eur(importi["spese_generali"])),
            ("Cassa Avvocati ( 4% )", _eur(importi["cpa"])),
            ("Totale imponibile", _eur(importi["imponibile"])),
            ("IVA 22% su Imponibile", _eur(importi["iva"])),
            ("IPOTESI DI COMPENSO LIQUIDABILE", _eur(importi["liquidabile"])),
            ("A dedurre ritenuta d'acconto 20% (su compenso e spese imponibili)", _eur(importi["ritenuta"])),
            ("Totale documento", _eur(importi["totale_documento"])),
        ]

    def _in_grassetto(voce: str) -> bool:
        return (
            voce in ("Fase", "PROSPETTO FINALE", "Totale documento")
            or voce.startswith("Compenso tabellare (")
            or voce.startswith("IPOTESI")
        )

    tabella = doc.add_table(rows=0, cols=2)
    for voce, importo in righe:
        celle = tabella.add_row().cells
        celle[0].text = voce
        celle[1].text = importo
        if _in_grassetto(voce):
            for cella in celle:
                for parag in cella.paragraphs:
                    for run in parag.runs:
                        run.bold = True

    # ---- oneri accessori ----
    if tipo == "monitorio":
        cu = _d2(contributo_unificato) if contributo_unificato >= 0 else _cu_monitorio(valore_causa)
        marca = Decimal("27.00")
        totale_complessivo = _d2(importi["liquidabile"] + cu + marca)
        par(
            "Nel totale sopra indicato non sono compresi gli oneri accessori che, per completezza e "
            "trasparenza, Vi trasmettiamo di seguito: al compenso liquidabile sopra indicato andranno "
            f"aggiunti {_eur(cu)} a titolo di contributo unificato ed {_eur(marca)} per la marca da "
            f"bollo, per un totale complessivo preventivato di {_eur(totale_complessivo)}."
        )
    elif tipo == "esecuzione":
        cu = _d2(contributo_unificato) if contributo_unificato >= 0 else Decimal("139.00")
        marca = Decimal("27.00")
        forfait = Decimal("120.00")
        totale_complessivo = _d2(importi["liquidabile"] + cu + marca + forfait)
        par(
            "Nel totale sopra indicato non sono compresi gli oneri accessori che, per completezza e "
            f"trasparenza, Vi trasmettiamo di seguito: al valore indicato andranno aggiunti {_eur(cu)} "
            f"a titolo di contributo unificato, {_eur(marca)} per la marca da bollo ed {_eur(forfait)} "
            "a titolo forfettario per le spese di deposito e postali, per un totale complessivo "
            f"preventivato di {_eur(totale_complessivo)}."
        )
        par(
            "Con riferimento al costo di deposito dell’atto di pignoramento presso lo sportello del "
            "competente Tribunale e al costo della trasferta andranno poi aggiunti € 200,00. A tali "
            "costi andranno poi sommate le spese delle visure camerali aggiornate, da € 10,00 a "
            "€ 15,00, ed eventuali spese di trasferte per l’udienza."
        )
    else:  # opposizione
        par(
            "Nel prospetto che precede non è compreso il contributo unificato, che nel giudizio di "
            "opposizione è a carico della parte opponente, oltre € 27,00 per la marca da bollo ed "
            "eventuali spese vive di trasferta."
        )

    if tipo in ("monitorio", "opposizione"):
        par(
            "Con riferimento all’imposta di registro, si precisa che la stessa sarà dovuta nella "
            "misura fissa di € 200,00 qualora le somme oggetto di condanna risultino soggette ad IVA, "
            "ovvero nella misura proporzionale del 3% sul valore della controversia in tutte le altre "
            "ipotesi."
        )

    par("Restiamo a Vostra disposizione per qualsiasi chiarimento o approfondimento.")
    par("Cordiali saluti,", align=WD_ALIGN_PARAGRAPH.LEFT, dopo=14)

    firme = doc.add_table(rows=1, cols=max(2, len(difensori)))
    for indice, nome in enumerate(difensori):
        firme.cell(0, indice).text = nome

    doc.add_paragraph()
    doc.add_paragraph()
    par(
        "Per integrale accettazione della presente quotazione e del relativo preventivo di spesa:",
        align=WD_ALIGN_PARAGRAPH.LEFT,
        dopo=18,
    )
    par(accettazione_denominazione or cliente_denominazione, bold=True, align=WD_ALIGN_PARAGRAPH.LEFT)
    par("_________________________________________", align=WD_ALIGN_PARAGRAPH.LEFT, dopo=0)

    return _salva(doc, f"quotazione_{tipo}_{debitore}")
