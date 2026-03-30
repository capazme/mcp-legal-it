"""Catalogo e dispatcher per 100 modelli di atti legali italiani.

Fornisce genera_modello_atto() che restituisce i metadati necessari per comporre
qualsiasi tipo di atto: struttura, campi, tool di calcolo da chiamare, resource
modello da leggere, riferimenti normativi e avvertenze.
"""

import json
from pathlib import Path

from src.server import mcp

_DATA = Path(__file__).resolve().parent.parent / "data"

with open(_DATA / "modelli_atti.json") as f:
    _CATALOGO: dict[str, dict] = json.load(f)

# Reverse index: map tool names and keywords to tipo_atto entries
_CATEGORIE = sorted({v["categoria"] for v in _CATALOGO.values()})


@mcp.tool(tags={"atti"})
def genera_modello_atto(tipo_atto: str, parametri: dict | None = None) -> dict:
    """Restituisce i metadati per comporre un atto legale: struttura, campi obbligatori,
    tool di calcolo da chiamare, resource modello da leggere, e riferimenti normativi.

    Chiamare PRIMA di redigere un atto per conoscere struttura e requisiti.
    Per l'elenco completo dei tipi disponibili, chiamare con tipo_atto="catalogo".
    Per cercare per categoria, chiamare con tipo_atto="cerca" e parametri={"query": "termine"}.

    Args:
        tipo_atto: Identificativo del tipo di atto. Usare "catalogo" per l'elenco completo,
                   "cerca" per cercare per parola chiave.
        parametri: Parametri opzionali dell'atto (per validazione campi obbligatori) o
                   {"query": "termine"} quando tipo_atto="cerca"
    """
    if parametri is None:
        parametri = {}

    # Catalog mode: return all available types grouped by category
    if tipo_atto == "catalogo":
        catalogo_per_categoria: dict[str, list[dict]] = {}
        for k, v in _CATALOGO.items():
            cat = v["categoria"]
            if cat not in catalogo_per_categoria:
                catalogo_per_categoria[cat] = []
            catalogo_per_categoria[cat].append({
                "tipo_atto": k,
                "descrizione": v["descrizione"],
                "tier": v["tier"],
            })
        return {
            "totale_tipi": len(_CATALOGO),
            "categorie": _CATEGORIE,
            "catalogo": catalogo_per_categoria,
        }

    # Search mode: find types matching a keyword
    if tipo_atto == "cerca":
        query = parametri.get("query", "").lower()
        if not query:
            return {"errore": "Specificare parametri={'query': 'termine'} per la ricerca"}
        risultati = []
        for k, v in _CATALOGO.items():
            searchable = f"{k} {v['descrizione']} {v['categoria']} {' '.join(v.get('riferimenti_normativi', []))}".lower()
            if query in searchable:
                risultati.append({
                    "tipo_atto": k,
                    "descrizione": v["descrizione"],
                    "categoria": v["categoria"],
                    "tier": v["tier"],
                })
        return {
            "query": query,
            "risultati": risultati,
            "totale": len(risultati),
        }

    # Lookup specific type
    entry = _CATALOGO.get(tipo_atto)
    if entry is None:
        # Fuzzy match: try substring matching
        candidati = [
            {"tipo_atto": k, "descrizione": v["descrizione"]}
            for k, v in _CATALOGO.items()
            if tipo_atto.lower() in k.lower() or tipo_atto.lower() in v["descrizione"].lower()
        ]
        return {
            "errore": f"Tipo atto '{tipo_atto}' non trovato nel catalogo",
            "suggerimenti": candidati[:10] if candidati else [],
            "nota": "Usare tipo_atto='catalogo' per l'elenco completo",
        }

    routing = entry["routing"]
    campi_obbligatori = entry["campi_obbligatori"]

    # Validate provided parameters against required fields
    campi_mancanti = [c for c in campi_obbligatori if c not in parametri] if parametri else campi_obbligatori

    # Build routing info
    result = {
        "tipo_atto": tipo_atto,
        "categoria": entry["categoria"],
        "descrizione": entry["descrizione"],
        "campi_obbligatori": campi_obbligatori,
        "campi_opzionali": entry.get("campi_opzionali", []),
        "campi_mancanti": campi_mancanti,
        "tool_calcolo": entry.get("tool_calcolo", []),
        "riferimenti_normativi": entry.get("riferimenti_normativi", []),
        "avvertenze": entry.get("avvertenze", []),
    }

    # Routing-specific info
    if routing["tipo"] == "tool_diretto":
        result["tool_diretto"] = routing["tool"]
        result["parametri_fissi"] = routing.get("parametri_fissi", {})
        result["istruzioni"] = f"Chiamare il tool `{routing['tool']}` con i parametri indicati."
    elif routing["tipo"] == "tool_enhance":
        result["tool_diretto"] = routing["tool"]
        result["parametri_fissi"] = routing.get("parametri_fissi", {})
        fase = routing.get("fase", 2)
        result["istruzioni"] = (
            f"Questo tipo di atto sarà supportato nella Fase {fase}. "
            f"Nel frattempo, usare il tool `{routing['tool']}` come base e adattare l'output."
        )
        result["disponibile_da_fase"] = fase
    elif routing["tipo"] == "resource":
        result["resource_modello"] = routing["resource"]
        result["istruzioni"] = (
            f"Leggere il modello dalla resource `{routing['resource']}`, "
            "compilare i campi con i dati forniti, e chiamare i tool di calcolo indicati."
        )
        fase = routing.get("fase", 3)
        result["disponibile_da_fase"] = fase
    elif routing["tipo"] == "preventivo_procedura":
        result["tool_diretto"] = "preventivo_procedura"
        result["parametri_fissi"] = routing.get("parametri_fissi", {})
        fase = routing.get("fase", 4)
        result["istruzioni"] = (
            f"Questo preventivo sarà supportato nella Fase {fase} via `preventivo_procedura()`. "
            "Nel frattempo, usare `preventivo_civile()` come approssimazione."
        )
        result["disponibile_da_fase"] = fase

    return result


@mcp.tool(tags={"atti"})
def esporta_atto_docx(
    testo: str,
    titolo: str = "Atto",
    autore: str = "",
) -> str:
    """Esporta un testo (atto, parere, bozza) in formato DOCX (Microsoft Word).

    Accetta testo in formato Markdown semplice e produce un file .docx formattato.
    Usare dopo aver generato un atto con genera_modello_atto() o un parere con il prompt parere_legale.
    Il file viene salvato nella directory temporanea e il percorso restituito.

    Args:
        testo: Testo dell'atto in formato Markdown (supporta: # titoli, **grassetto**, *corsivo*, - elenchi, paragrafi)
        titolo: Titolo del documento (usato come nome file e intestazione)
        autore: Nome dell'autore (opzionale, inserito nei metadati del documento)
    """
    try:
        from docx import Document
        from docx.opc.constants import RELATIONSHIP_TYPE as RT  # noqa: F401
    except ImportError:
        return "Errore: python-docx non installato. Eseguire: pip install python-docx"

    if not testo or not testo.strip():
        return "Errore: il testo dell'atto è vuoto."

    import re

    def _sanitize_filename(name: str) -> str:
        name = name.lower().replace(" ", "_")
        name = re.sub(r"[^a-z0-9_\-]", "", name)
        return name[:50] or "atto"

    def _add_formatted_paragraph(doc, text, style=None):
        if style:
            para = doc.add_paragraph(style=style)
        else:
            para = doc.add_paragraph()

        # Split on bold (**...**) and italic (*...*) markers
        # Pattern: **bold**, *italic* (bold checked first to avoid conflict)
        parts = re.split(r"(\*\*[^*]+\*\*|\*[^*]+\*)", text)
        for part in parts:
            if part.startswith("**") and part.endswith("**"):
                run = para.add_run(part[2:-2])
                run.bold = True
            elif part.startswith("*") and part.endswith("*"):
                run = para.add_run(part[1:-1])
                run.italic = True
            else:
                para.add_run(part)
        return para

    doc = Document()

    # Set core properties
    props = doc.core_properties
    props.title = titolo
    if autore:
        props.author = autore

    lines = testo.split("\n")
    for line in lines:
        stripped = line.rstrip()

        if stripped.startswith("### "):
            doc.add_heading(stripped[4:], level=3)
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:], level=2)
        elif stripped.startswith("# "):
            doc.add_heading(stripped[2:], level=1)
        elif re.match(r"^[-*] ", stripped):
            _add_formatted_paragraph(doc, stripped[2:], style="List Bullet")
        elif re.match(r"^\d+\. ", stripped):
            content = re.sub(r"^\d+\. ", "", stripped)
            _add_formatted_paragraph(doc, content, style="List Number")
        elif stripped.startswith("> "):
            para = _add_formatted_paragraph(doc, stripped[2:])
            for run in para.runs:
                run.italic = True
        elif stripped == "":
            doc.add_paragraph()
        else:
            _add_formatted_paragraph(doc, stripped)

    # Save to temp file
    import os
    tmp_dir = "/tmp/mcp-legal-it"
    os.makedirs(tmp_dir, exist_ok=True)
    filename = _sanitize_filename(titolo) + ".docx"
    filepath = os.path.join(tmp_dir, filename)
    doc.save(filepath)

    size_kb = round(os.path.getsize(filepath) / 1024, 1)
    return f"File salvato: {filepath} ({size_kb} KB)"


@mcp.tool(tags={"atti"})
def lista_categorie_atti() -> dict:
    """Restituisce le categorie di atti disponibili con il conteggio per ciascuna.
    Utile per orientare l'utente nella scelta del tipo di atto.
    """
    conteggio: dict[str, int] = {}
    for v in _CATALOGO.values():
        cat = v["categoria"]
        conteggio[cat] = conteggio.get(cat, 0) + 1
    return {
        "categorie": [
            {"nome": cat, "totale": conteggio[cat]}
            for cat in sorted(conteggio, key=conteggio.get, reverse=True)
        ],
        "totale_atti": len(_CATALOGO),
    }
