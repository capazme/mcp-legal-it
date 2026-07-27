"""Test per src/tools/procure_quotazioni.py — procura alle liti e quotazioni DOCX.

I valori attesi dei prospetti sono quelli validati su pratiche reali di recupero
crediti seriale (D.M. 55/2014 agg. D.M. 147/2022): monitorio scaglione fino a
€ 5.200 → totale documento € 378,70; scaglione € 5.201-26.000 → € 453,78;
opposizione a valori medi su € 55.898,69 → € 22.534,57; esecuzione forzata
(minimi, scaglione base) → liquidabile € 656,60 e complessivo € 942,60.
"""

import importlib
import os

_mod = importlib.import_module("src.tools.procure_quotazioni")


def _call(fn_name, **kwargs):
    fn = getattr(_mod, fn_name)
    fn = getattr(fn, "fn", fn)
    return fn(**kwargs)


def _leggi_docx(result: str) -> str:
    from docx import Document

    path = result.split("File salvato: ")[1].split(" (")[0]
    assert os.path.isfile(path)
    doc = Document(path)
    parti = [p.text for p in doc.paragraphs]
    for tabella in doc.tables:
        for riga in tabella.rows:
            parti.append(" | ".join(cella.text for cella in riga.cells))
    return "\n".join(parti)


_PROCURA_KWARGS = dict(
    mandante_denominazione="Esempio S.r.l.",
    mandante_sede="via Roma n. 1 – 20100 Milano (MI)",
    mandante_cf_piva="00000000000",
    firmatario_nome="Mario Rossi",
    firmatario_cf="RSSMRA70A01F205X",
    controparte=(
        "Delta S.r.l., in persona del proprio legale rappresentante pro tempore, "
        "con sede legale in via Verdi n. 2 – 20100 Milano (MI), Cod. Fiscale e "
        "Partita IVA 11111111111"
    ),
    difensori=[
        {"nome": "Giulia Bianchi", "cf": "BNCGLI80A41F205Y"},
        {"nome": "Paolo Verdi", "cf": "VRDPLA75B02F205Z"},
    ],
    domicilio_studio="corso Esempio n. 10 – 20100 Milano (MI)",
    pec="g.bianchi@pec.esempio.it e p.verdi@pec.esempio.it",
)

_QUOTAZIONE_KWARGS = dict(
    cliente_denominazione="Esempio S.r.l.",
    cliente_indirizzo="via Roma n. 1; 20100 - Milano",
    difensori=["Avv. Giulia Bianchi", "Avv. Paolo Verdi"],
)


class TestGeneraProcuraLitiDocx:

    def test_generazione_base(self):
        result = _call("genera_procura_liti_docx", **_PROCURA_KWARGS)
        assert "File salvato" in result
        assert "errore" not in result.lower()

    def test_contenuto_verbatim_e_struttura(self):
        result = _call("genera_procura_liti_docx", **_PROCURA_KWARGS, data_documento="22/07/2026")
        testo = _leggi_docx(result)
        assert "PROCURA ALLE LITI" in testo
        assert "ART. 83, COMMA 3, C.P.C." in testo
        # clausola controparte verbatim
        assert "Partita IVA 11111111111" in testo
        # conversione data GG/MM/AAAA -> forma estesa
        assert "22 luglio 2026" in testo
        # entrambi i difensori con autentica
        assert "Avv. Giulia Bianchi" in testo
        assert "Avv. Paolo Verdi" in testo
        assert "È vera e autentica" in testo
        assert "congiuntamente e disgiuntamente" in testo

    def test_qualifica_firmatario(self):
        kwargs = dict(_PROCURA_KWARGS)
        kwargs["firmatario_qualifica"] = "presidente del consiglio di amministrazione e legale rappresentante"
        testo = _leggi_docx(_call("genera_procura_liti_docx", **kwargs))
        assert "presidente del consiglio di amministrazione" in testo

    def test_difensore_singolo(self):
        kwargs = dict(_PROCURA_KWARGS)
        kwargs["difensori"] = [{"nome": "Giulia Bianchi", "cf": "BNCGLI80A41F205Y"}]
        testo = _leggi_docx(_call("genera_procura_liti_docx", **kwargs))
        assert "congiuntamente e disgiuntamente" not in testo
        assert "il nominato difensore" in testo

    def test_errore_senza_difensori(self):
        kwargs = dict(_PROCURA_KWARGS)
        kwargs["difensori"] = []
        result = _call("genera_procura_liti_docx", **kwargs)
        assert "errore" in result.lower()

    def test_errore_controparte_vuota(self):
        kwargs = dict(_PROCURA_KWARGS)
        kwargs["controparte"] = "  "
        result = _call("genera_procura_liti_docx", **kwargs)
        assert "errore" in result.lower()


class TestGeneraQuotazioneDocx:

    def test_monitorio_scaglione_base(self):
        result = _call(
            "genera_quotazione_docx",
            tipo="monitorio", valore_causa=3633.77, debitore="Delta S.r.l.",
            **_QUOTAZIONE_KWARGS,
        )
        testo = _leggi_docx(result)
        assert "€ 237,00" in testo          # fase unica, minimo
        assert "€ 378,70" in testo          # totale documento
        assert "€ 49,00" in testo           # CU dimezzato
        assert "€ 525,56" in testo          # complessivo con oneri
        assert "Per integrale accettazione" in testo

    def test_monitorio_secondo_scaglione(self):
        result = _call(
            "genera_quotazione_docx",
            tipo="monitorio", valore_causa=7428.04, debitore="Delta S.r.l.",
            **_QUOTAZIONE_KWARGS,
        )
        testo = _leggi_docx(result)
        assert "€ 284,00" in testo
        assert "€ 453,78" in testo
        assert "€ 118,50" in testo          # CU dimezzato scaglione 5.201-26.000

    def test_monitorio_medi_usa_tabella_ministeriale(self):
        result = _call(
            "genera_quotazione_docx",
            tipo="monitorio", valore_causa=3000, debitore="Delta S.r.l.",
            livello="medi", **_QUOTAZIONE_KWARGS,
        )
        testo = _leggi_docx(result)
        assert "valore medio" in testo
        assert "€ 473,00" in testo          # medio ministeriale (non 237 x 2 = 474)

    def test_monitorio_scaglione_esteso_520k(self):
        result = _call(
            "genera_quotazione_docx",
            tipo="monitorio", valore_causa=300000, debitore="Delta S.r.l.",
            **_QUOTAZIONE_KWARGS,
        )
        testo = _leggi_docx(result)
        assert "€ 2.197,00" in testo        # minimo scaglione 260.001-520.000
        assert "€ 607,00" in testo          # CU dalla tabella canonica

    def test_opposizione_medi(self):
        result = _call(
            "genera_quotazione_docx",
            tipo="opposizione", valore_causa=55898.69, debitore="Gamma S.r.l.",
            livello="medi", **_QUOTAZIONE_KWARGS,
        )
        testo = _leggi_docx(result)
        assert "art. 645 c.p.c." in testo
        assert "€ 14.103,00" in testo       # somma 4 fasi valori medi
        assert "€ 22.534,57" in testo       # totale documento
        assert "a carico della parte opponente" in testo

    def test_opposizione_minimi_scaglione_basso(self):
        result = _call(
            "genera_quotazione_docx",
            tipo="opposizione", valore_causa=918.29, debitore="Gamma S.r.l.",
            **_QUOTAZIONE_KWARGS,
        )
        testo = _leggi_docx(result)
        assert "€ 332,00" in testo          # 66+66+100+100 minimi fino a 1.100
        assert "€ 530,48" in testo          # totale documento

    def test_esecuzione(self):
        result = _call(
            "genera_quotazione_docx",
            tipo="esecuzione", valore_causa=1937.67, debitore="Delta S.r.l.",
            **_QUOTAZIONE_KWARGS,
        )
        testo = _leggi_docx(result)
        assert "€ 450,00" in testo          # 166 + 284
        assert "€ 656,60" in testo          # liquidabile
        assert "€ 942,60" in testo          # con CU 139 + marca 27 + forfait 120
        assert "pignoramento" in testo
        assert "ritenuta" not in testo.lower()   # il prospetto esecuzione non ha RA
        assert "PCT" not in testo                # e nessun aumento 30%

    def test_accettazione_personalizzata_senza_persona(self):
        result = _call(
            "genera_quotazione_docx",
            tipo="monitorio", valore_causa=1000, debitore="Delta S.r.l.",
            accettazione_denominazione="Esempio S.r.l. (Brand)",
            **_QUOTAZIONE_KWARGS,
        )
        testo = _leggi_docx(result)
        assert "Esempio S.r.l. (Brand)" in testo
        assert "Mario Rossi" not in testo   # nel blocco accettazione niente persone fisiche

    def test_cu_override(self):
        result = _call(
            "genera_quotazione_docx",
            tipo="monitorio", valore_causa=1000, debitore="Delta S.r.l.",
            contributo_unificato=21.50, **_QUOTAZIONE_KWARGS,
        )
        testo = _leggi_docx(result)
        assert "€ 21,50" in testo

    def test_errori_input(self):
        base = dict(valore_causa=1000, debitore="X", **_QUOTAZIONE_KWARGS)
        assert "errore" in _call("genera_quotazione_docx", tipo="cautelare", **base).lower()
        assert "errore" in _call(
            "genera_quotazione_docx", tipo="monitorio", livello="massimi",
            valore_causa=1000, debitore="X", **_QUOTAZIONE_KWARGS,
        ).lower()
        assert "errore" in _call(
            "genera_quotazione_docx", tipo="monitorio",
            valore_causa=0, debitore="X", **_QUOTAZIONE_KWARGS,
        ).lower()
        assert "errore" in _call(
            "genera_quotazione_docx", tipo="monitorio",
            valore_causa=600000, debitore="X", **_QUOTAZIONE_KWARGS,
        ).lower()

    def test_esecuzione_guardie_default(self):
        # default (minimi scaglione base) + livello medi -> errore, non etichetta falsa
        r = _call(
            "genera_quotazione_docx", tipo="esecuzione",
            valore_causa=3000, debitore="X", livello="medi", **_QUOTAZIONE_KWARGS,
        )
        assert "errore" in r.lower()
        # default + valore oltre lo scaglione base -> errore, non quotazione fuori scaglione
        r = _call(
            "genera_quotazione_docx", tipo="esecuzione",
            valore_causa=50000, debitore="X", **_QUOTAZIONE_KWARGS,
        )
        assert "errore" in r.lower()
        # compensi negativi -> errore
        r = _call(
            "genera_quotazione_docx", tipo="esecuzione",
            valore_causa=3000, debitore="X",
            compenso_fase_introduttiva=-166, compenso_fase_trattazione=-284,
            **_QUOTAZIONE_KWARGS,
        )
        assert "errore" in r.lower()
        # override esplicito per scaglione alto -> funziona
        r = _call(
            "genera_quotazione_docx", tipo="esecuzione",
            valore_causa=50000, debitore="X",
            compenso_fase_introduttiva=482, compenso_fase_trattazione=925,
            **_QUOTAZIONE_KWARGS,
        )
        assert "File salvato" in r
